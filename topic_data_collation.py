# build_labels_and_splits.py
# -*- coding: utf-8 -*-

import re
import json
from pathlib import Path

import pandas as pd
from sklearn.model_selection import train_test_split

try:
    from docx import Document  # python-docx
except Exception:
    Document = None  # 若未安裝，稍後會用 fallback 對應表


# ====== 路徑設定 ======
TOPIC_PROP_CSV = "output_bertopic\\topic_proportions.csv"        # 含 image_name、topic_0~topicN、dominant_topic、dominant_prob
TOPIC_META_CSV = "output_bertopic\\topic.csv"                    # 用不到也可不放，但保留參考
DOCX_PATH      = "小主題聚合成大主題_0930.docx"   # 專家合併後的 12 大主題對應來源
OUT_LABELS_CSV = "bertopic_labels.csv"
OUT_SPLIT_CSV  = "split.csv"

# ====== 亂數種子（確保重現性） ======
SEED = 42

# ====== 後備對應表：{小主題topic_id -> 新的大主題編號(1~12)} ======
# 取自 docx（若 docx 解析失敗會使用此表；若成功解析則以 docx 為準）
FALLBACK_MAJOR_TO_SMALL = {
    1:  [0, 36, 55, 62, 64, 67, 98, 101],
    2:  [8, 19, 20, 41, 65, 83, 90, 97, 99],
    3:  [6, 10, 28, 59, 96, 100, 103, 107],   # 注意：100 同時出現在 1 與 3，程式會警告並以先出現者為主
    4:  [2, 17, 26, 29, 38, 81, 91, 102],
    5:  [33, 44, 46, 50, 53, 56, 76, 80, 84],
    6:  [35, 48, 71, 72, 85, 89, 92, 94],
    7:  [22, 30, 32, 60, 69, 70, 77, 95],
    8:  [1, 7, 14, 18, 24, 45, 52, 54, 73],
    9:  [3, 11, 13, 57, 82, 86, 104],
    10: [5, 12, 40, 43, 49, 58, 61, 63, 78],
    11: [4, 15, 21, 23, 25, 51, 87, 88, 105],
    12: [9, 31, 37, 39, 42, 68, 74, 75, 79],
}

def build_mapping_from_dict(major_to_small: dict) -> dict:
    """將 {大主題: [小主題...]} 轉為 {小主題: 大主題}；若小主題重複歸屬，保留「先出現」者並警告。"""
    small_to_major = {}
    duplicates = {}
    for major, small_list in major_to_small.items():
        for s in small_list:
            if s in small_to_major and small_to_major[s] != major:
                duplicates.setdefault(s, set()).update({small_to_major[s], major})
                # 保留先出現者，不覆蓋
            else:
                small_to_major[s] = major
    if duplicates:
        print("[WARN] 發現重複歸屬的小主題：", json.dumps({k: sorted(list(v)) for k,v in duplicates.items()}, ensure_ascii=False))
        print("[WARN] 已採用『先出現者優先』，未改動先前的歸屬。")
    return small_to_major


def try_parse_docx_mapping(docx_path: str) -> dict | None:
    """嘗試從 docx 解析 {小主題topic_id -> 大主題編號}；失敗回傳 None。"""
    if Document is None:
        print("[INFO] 未安裝 python-docx 或導入失敗，改用後備對應表。")
        return None
    p = Path(docx_path)
    if not p.exists():
        print(f"[INFO] 找不到 docx：{docx_path}，改用後備對應表。")
        return None

    doc = Document(docx_path)

    # 解析策略：
    # 1) 優先巡覽 tables：很多整理檔會用表格，找出一列中含「小主題編號」清單（逗號分隔的數字）。
    # 2) 若 tables 無法解析，再掃描 paragraphs：遇到「^\d+\.\s」視為大主題編號，下一段若是數字逗號清單就解析。
    major_to_small = {}

    # --- 試表格 ---
    numeric_line_re = re.compile(r"^\s*\d+(?:\s*,\s*\d+)*\s*$")
    found_any = False

    for t in doc.tables:
        # 嘗試逐列找出「大主題標題 + 小主題編號」的成對資訊
        current_major = None
        for row in t.rows:
            cells_text = [c.text.strip() for c in row.cells]
            # 嘗試在某格找出像 "1. 現代都市..." 的大主題編號
            for txt in cells_text:
                m = re.match(r"^\s*(\d+)\.\s", txt)
                if m:
                    current_major = int(m.group(1))
                    if current_major not in major_to_small:
                        major_to_small[current_major] = []
            # 嘗試抓小主題清單
            for txt in cells_text:
                if numeric_line_re.match(txt) and current_major is not None:
                    ids = [int(x) for x in re.findall(r"\d+", txt)]
                    major_to_small[current_major].extend(ids)
                    found_any = True

    # --- 試段落 ---
    if not found_any:
        current_major = None
        prev_was_major = False
        for para in doc.paragraphs:
            text = para.text.strip()
            # 大主題標題行
            m = re.match(r"^\s*(\d+)\.\s", text)
            if m:
                current_major = int(m.group(1))
                major_to_small.setdefault(current_major, [])
                prev_was_major = True
                continue

            # 下一行若是「數字逗號清單」，視為該大主題的小主題編號
            if prev_was_major and numeric_line_re.match(text) and current_major is not None:
                ids = [int(x) for x in re.findall(r"\d+", text)]
                major_to_small[current_major].extend(ids)
                found_any = True
                prev_was_major = False
            else:
                prev_was_major = False  # reset

    if not found_any:
        print("[INFO] 無法自動從 docx 解析小主題→大主題，改用後備對應表。")
        return None

    # 清理：去重並排序
    for k in list(major_to_small.keys()):
        major_to_small[k] = sorted(set(major_to_small[k]))

    # 回傳小→大映射
    return build_mapping_from_dict(major_to_small)


def ensure_int_topic(x):
    """容許 'topic_23'、'23'、23、'-1' 等表示法，回傳 int 或 None。"""
    if pd.isna(x):
        return None
    m = re.search(r"-?\d+", str(x))
    return int(m.group(0)) if m else None


def main():
    # 1) 建立 小主題(topic_id) -> 大主題(1~12) 對應
    small_to_major = try_parse_docx_mapping(DOCX_PATH)
    if small_to_major is None:
        small_to_major = build_mapping_from_dict(FALLBACK_MAJOR_TO_SMALL)

    # 2) 讀取 topic_proportions.csv
    df = pd.read_csv(TOPIC_PROP_CSV)

    # 2.1 確保 dominant_topic / dominant_prob 存在；若 dominant_topic 缺失則由 topic_欄位計算
    topic_cols = [c for c in df.columns if re.fullmatch(r"topic_-?\d+|topic_\d+", c) or re.fullmatch(r"topic_\d+", c)]
    if not topic_cols:
        # 也支援 "topic_0~topic106" 的簡單命名（不含負號）
        topic_cols = [c for c in df.columns if c.startswith("topic_")]

    if "dominant_topic" not in df.columns and topic_cols:
        # 取 row 中機率最高的 topic_*
        max_idx = df[topic_cols].astype(float).idxmax(axis=1)  # e.g. 'topic_23'
        df["dominant_topic"] = max_idx.str.extract(r"(-?\d+)").astype(int)

    # 2.2 形成輸出欄位
    labels = pd.DataFrame()
    labels["image_name"] = df["image_name"]

    # topic_raw：對應原 BERTopic 輸出的 topic id（等同 dominant_topic 的整數化）
    if "dominant_topic" in df.columns:
        labels["topic_raw"] = df["dominant_topic"].apply(ensure_int_topic)
        labels["dominant_topic"] = df["dominant_topic"].apply(ensure_int_topic)
    else:
        raise ValueError("找不到 dominant_topic，也無法從 topic_* 欄自動計算。請確認輸入檔。")

    # dominant_prob
    if "dominant_prob" in df.columns:
        labels["dominant_prob"] = pd.to_numeric(df["dominant_prob"], errors="coerce")
    else:
        # 若未提供，則由最高機率欄位取得
        if topic_cols:
            labels["dominant_prob"] = df[topic_cols].astype(float).max(axis=1)
        else:
            labels["dominant_prob"] = None

    # new_topic_num：由小主題 topic id 映射到 docx 的 12 大主題編號；找不到給 -1
    labels["new_topic_num"] = labels["topic_raw"].map(small_to_major).fillna(-1).astype(int)

    # 3) 輸出 bertopic_labels.csv
    labels = labels[["image_name", "topic_raw", "dominant_topic", "dominant_prob", "new_topic_num"]]
    labels.to_csv(OUT_LABELS_CSV, index=False, encoding="utf-8-sig")
    print(f"[OK] 已輸出：{OUT_LABELS_CSV}  （共 {len(labels):,} 筆）")

    # 4) 依 6.4 : 1.6 : 2 產生 split.csv（train/val/test）
    #    先切出 test=20%，剩下 80% 再切出 val=20%/80%=0.25 → 0.16 總體
    unique_imgs = labels["image_name"].drop_duplicates()
    train_imgs, temp_imgs = train_test_split(unique_imgs, test_size=0.20, random_state=SEED, shuffle=True)
    val_imgs, test_imgs = train_test_split(temp_imgs, test_size=(20/36), random_state=SEED, shuffle=True)
    # 檢查比例
    n = len(unique_imgs)
    print(f"[INFO] 影像總數 = {n:,}")
    print(f"[INFO] train = {len(train_imgs):,} ({len(train_imgs)/n:.2%})")
    print(f"[INFO] val   = {len(val_imgs):,} ({len(val_imgs)/n:.2%})")
    print(f"[INFO] test  = {len(test_imgs):,} ({len(test_imgs)/n:.2%})")

    split_df = pd.concat([
        pd.DataFrame({"image_name": train_imgs, "split": "train"}),
        pd.DataFrame({"image_name": val_imgs,   "split": "val"}),
        pd.DataFrame({"image_name": test_imgs,  "split": "test"}),
    ], ignore_index=True)

    split_df.to_csv(OUT_SPLIT_CSV, index=False, encoding="utf-8-sig")
    print(f"[OK] 已輸出：{OUT_SPLIT_CSV}  （共 {len(split_df):,} 筆）")

    # 5) 友善提示：列出沒有對應到 12 大主題的 topic id（-1）
    unmapped = labels.loc[labels["new_topic_num"] == -1, "topic_raw"].dropna().unique()
    if len(unmapped) > 0:
        print("[WARN] 有 topic id 未對應到任何大主題（new_topic_num = -1）：", sorted(map(int, unmapped)))


if __name__ == "__main__":
    main()
